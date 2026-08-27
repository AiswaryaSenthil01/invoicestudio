"""
Word (.docx) Invoice Generator using python-docx.
Replicates the structured GST Tax Invoice layout.
"""
from pathlib import Path
from typing import Dict, Any, List
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from src.utils.calculations import split_rs_paise
from src.utils.number_to_words import amount_to_words
from src.utils.file_utils import get_invoice_storage_path


def set_cell_border(cell, **kwargs):
    """
    Set cell borders for docx table cells.
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='000000')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
    for border_name, border_style in kwargs.items():
        border_el = parse_xml(
            r'<w:%s %s w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            % (border_name, nsdecls('w'), border_style.get('val', 'single'),
               border_style.get('sz', '4'), border_style.get('color', '000000'))
        )
        tcBorders.append(border_el)
    tcPr.append(tcBorders)


class DocxGenerator:
    def generate(self, invoice_data: Dict[str, Any], output_path: Path | str = None) -> Path:
        """
        Generate Word (.docx) export for the invoice.
        """
        invoice_no = invoice_data.get("invoice_no", "INV-001")
        invoice_date = invoice_data.get("invoice_date", "")

        if output_path is None:
            output_path = get_invoice_storage_path(invoice_no, invoice_date, "docx")
        else:
            output_path = Path(output_path)

        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = docx.Document()

        # Set A4 Page and Margins (0.4 inch)
        for section in doc.sections:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

        # Style normal text
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(0, 0, 0)

        # 1. Header Grid Table (GSTIN/PAN left, Title/Company center, Copy Type right)
        hdr_tbl = doc.add_table(rows=1, cols=3)
        hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_tbl.autofit = False

        c_left, c_mid, c_right = hdr_tbl.rows[0].cells
        c_left.width = Inches(2.2)
        c_mid.width = Inches(3.6)
        c_right.width = Inches(1.6)

        # Top Left
        p_left = c_left.paragraphs[0]
        p_left.paragraph_format.space_after = Pt(2)
        r = p_left.add_run(f"GSTIN NO: {invoice_data.get('company_gstin', '33BKXPS7582P1ZR')}\n")
        r.bold = True
        r = p_left.add_run(f"PAN NO: {invoice_data.get('company_pan', 'BKXPS7582P')}")
        r.bold = True

        # Center Company Details
        p_mid = c_mid.paragraphs[0]
        p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mid.paragraph_format.space_after = Pt(2)
        
        r_title = p_mid.add_run("TAX INVOICE\n")
        r_title.bold = True
        r_title.font.size = Pt(13)

        r_comp = p_mid.add_run(f"{invoice_data.get('company_name', 'NAMURA ENGG. WORKS')}\n")
        r_comp.bold = True
        r_comp.font.size = Pt(11)

        p_mid.add_run(f"{invoice_data.get('company_address', '')}\n")
        p_mid.add_run(f"Phone No: {invoice_data.get('company_phone', '')}   State Code: {invoice_data.get('company_state_code', '')}\n")
        p_mid.add_run(f"E-mail: {invoice_data.get('company_email', '')}")

        # Top Right
        p_right = c_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_copy = p_right.add_run(invoice_data.get('copy_type', 'Original Copy'))
        r_copy.italic = True
        r_copy.bold = True

        # 2. Meta & Party Details Table
        meta_tbl = doc.add_table(rows=2, cols=2)
        meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_tbl.autofit = False

        for row in meta_tbl.rows:
            row.cells[0].width = Inches(3.7)
            row.cells[1].width = Inches(3.7)
            for cell in row.cells:
                set_cell_border(cell, top=dict(sz=4, val='single', color='333333'),
                                      bottom=dict(sz=4, val='single', color='333333'),
                                      left=dict(sz=4, val='single', color='333333'),
                                      right=dict(sz=4, val='single', color='333333'))

        # Row 0: Invoice Meta
        c_inv_l = meta_tbl.rows[0].cells[0]
        p = c_inv_l.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"Invoice No : {invoice_data.get('invoice_no', '')}\n").bold = True
        p.add_run(f"Date       : {invoice_data.get('invoice_date', '')}")

        c_inv_r = meta_tbl.rows[0].cells[1]
        p = c_inv_r.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"P.O. No    : {invoice_data.get('po_no', '')}\n")
        p.add_run(f"P.O. Date  : {invoice_data.get('po_date', '')}\n")
        p.add_run(f"Vehicle No : {invoice_data.get('vehicle_no', '')}")

        # Row 1: Billed To & Shipped To
        c_bill = meta_tbl.rows[1].cells[0]
        p = c_bill.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.add_run("Billed To:\n").bold = True
        p.add_run(f"{invoice_data.get('billed_to_name', '')}\n").bold = True
        p.add_run(f"{invoice_data.get('billed_to_address', '')}\n")
        extra_b = []
        if invoice_data.get("billed_to_phone"):
            extra_b.append(f"Ph: {invoice_data.get('billed_to_phone')}")
        if invoice_data.get("billed_to_gstin"):
            extra_b.append(f"GSTIN: {invoice_data.get('billed_to_gstin')}")
        if invoice_data.get("billed_to_state_code"):
            extra_b.append(f"State: {invoice_data.get('billed_to_state_code')}")
        if extra_b:
            p.add_run(" | ".join(extra_b))

        c_ship = meta_tbl.rows[1].cells[1]
        p = c_ship.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.add_run("Shipped To:\n").bold = True
        ship_name = invoice_data.get('shipped_to_name') or invoice_data.get('billed_to_name', '')
        p.add_run(f"{ship_name}\n").bold = True
        ship_addr = invoice_data.get('shipped_to_address') or invoice_data.get('billed_to_address', '')
        p.add_run(f"{ship_addr}\n")
        extra_s = []
        if invoice_data.get("shipped_to_phone"):
            extra_s.append(f"Ph: {invoice_data.get('shipped_to_phone')}")
        if invoice_data.get("shipped_to_gstin"):
            extra_s.append(f"GSTIN: {invoice_data.get('shipped_to_gstin')}")
        if invoice_data.get("shipped_to_state_code"):
            extra_s.append(f"State: {invoice_data.get('shipped_to_state_code')}")
        if extra_s:
            p.add_run(" | ".join(extra_s))

        # 3. Items Table
        items = invoice_data.get("items", [])
        # Columns: S.NO (0.5"), Description (2.9"), HSN (0.8"), QTY (0.6"), RATE Rs (0.9"), RATE P (0.4"), AMT Rs (0.9"), AMT P (0.4")
        col_widths = [Inches(0.5), Inches(2.9), Inches(0.8), Inches(0.6), Inches(0.9), Inches(0.4), Inches(0.9), Inches(0.4)]
        
        # We create a table with 2 header rows + len(items) + 6 totals rows
        item_tbl = doc.add_table(rows=2 + len(items) + 6, cols=8)
        item_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        item_tbl.autofit = False

        # Apply borders & widths to all cells
        for row in item_tbl.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = col_widths[idx]
                set_cell_border(cell, top=dict(sz=4, val='single', color='333333'),
                                      bottom=dict(sz=4, val='single', color='333333'),
                                      left=dict(sz=4, val='single', color='333333'),
                                      right=dict(sz=4, val='single', color='333333'))

        # Header Row 0
        h_row0 = item_tbl.rows[0]
        h_row1 = item_tbl.rows[1]

        h_row0.cells[0].paragraphs[0].add_run("S.NO").bold = True
        h_row0.cells[1].paragraphs[0].add_run("Description of Goods").bold = True
        h_row0.cells[2].paragraphs[0].add_run("HSN CODE").bold = True
        h_row0.cells[3].paragraphs[0].add_run("QTY.").bold = True
        h_row0.cells[4].paragraphs[0].add_run("RATE").bold = True
        h_row0.cells[6].paragraphs[0].add_run("AMOUNT").bold = True

        h_row1.cells[4].paragraphs[0].add_run("Rs.")
        h_row1.cells[5].paragraphs[0].add_run("P")
        h_row1.cells[6].paragraphs[0].add_run("Rs.")
        h_row1.cells[7].paragraphs[0].add_run("P")

        # Line Items Rows
        for i, item in enumerate(items):
            row = item_tbl.rows[2 + i]
            row.cells[0].paragraphs[0].add_run(str(item.get("s_no", i + 1)))
            row.cells[1].paragraphs[0].add_run(str(item.get("description", "")))
            row.cells[2].paragraphs[0].add_run(str(item.get("hsn_code", "")))
            row.cells[3].paragraphs[0].add_run(str(item.get("quantity", "")))
            
            rate_rs, rate_p = split_rs_paise(item.get("rate", 0))
            amt_rs, amt_p = split_rs_paise(item.get("amount", 0))

            p = row.cells[4].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(rate_rs)

            p = row.cells[5].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(rate_p)

            p = row.cells[6].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(amt_rs)

            p = row.cells[7].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(amt_p)

        # Totals Rows
        base_tot_idx = 2 + len(items)
        
        # In words in left column (merge or place in S.NO / Desc)
        grand_total = invoice_data.get("grand_total", 0.0)
        in_words = invoice_data.get("total_in_words") or amount_to_words(grand_total)

        words_cell = item_tbl.rows[base_tot_idx].cells[0]
        p = words_cell.paragraphs[0]
        p.add_run("RUPEES:\n").bold = True
        r_w = p.add_run(in_words)
        r_w.italic = True

        totals_data = [
            ("TOTAL", invoice_data.get("subtotal", 0.0)),
            (f"CGST. {invoice_data.get('cgst_rate', 0.0):g}%", invoice_data.get("cgst_amount", 0.0) if invoice_data.get("cgst_rate", 0) > 0 else None),
            (f"SGST. {invoice_data.get('sgst_rate', 0.0):g}%", invoice_data.get("sgst_amount", 0.0) if invoice_data.get("sgst_rate", 0) > 0 else None),
            (f"IGST. {invoice_data.get('igst_rate', 0.0):g}%" if invoice_data.get("igst_rate", 0) > 0 else "IGST.", invoice_data.get("igst_amount", 0.0) if invoice_data.get("igst_rate", 0) > 0 else None),
            ("TOTAL TAX", invoice_data.get("total_tax", 0.0)),
            ("INVOICE TOTAL", grand_total)
        ]

        for offset, (lbl, val) in enumerate(totals_data):
            row = item_tbl.rows[base_tot_idx + offset]
            # Label in cell 4
            p_lbl = row.cells[4].paragraphs[0]
            r_lbl = p_lbl.add_run(lbl)
            if offset == 5:
                r_lbl.bold = True

            # Amount in cells 6 and 7
            rs_val, p_val = split_rs_paise(val)
            p_rs = row.cells[6].paragraphs[0]
            p_rs.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_rs = p_rs.add_run(rs_val)
            if offset == 5:
                r_rs.bold = True

            p_p = row.cells[7].paragraphs[0]
            p_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_p = p_p.add_run(p_val)
            if offset == 5:
                r_p.bold = True

        # 4. Footer Section
        doc.add_paragraph().paragraph_format.space_before = Pt(8)
        p_dec = doc.add_paragraph()
        p_dec.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_dec.paragraph_format.space_after = Pt(2)
        r_d = p_dec.add_run("Certified that the above particulars are true & correct\n")
        r_d.font.size = Pt(8)
        r_c = p_dec.add_run(f"For {invoice_data.get('company_name', 'NAMURA ENGG. WORKS')}")
        r_c.bold = True
        r_c.font.size = Pt(8.5)

        p_sig = doc.add_paragraph()
        p_sig.paragraph_format.space_before = Pt(30)
        p_sig.add_run("Receiver’s Signature                                                                            Authorized Signatory").bold = True

        doc.save(str(output_path))
        return output_path
